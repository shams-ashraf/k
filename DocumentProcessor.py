import streamlit as st
import re
import fitz
import glob
import os
import pickle
import hashlib
import docx
from dotenv import load_dotenv
from deep_translator import GoogleTranslator
from langdetect import detect, DetectorFactory
import langid

load_dotenv()

DetectorFactory.seed = 0

PDF_PASSWORD = os.getenv("PDF_PASSWORD", "")
DOCS_FOLDER = "/mount/src/k/documents"
CACHE_FOLDER = os.getenv("CACHE_FOLDER", "./cache")
TESSDATA_PATH = os.getenv("TESSDATA_PATH")

os.makedirs(DOCS_FOLDER, exist_ok=True)
os.makedirs(CACHE_FOLDER, exist_ok=True)

def detect_doc_language(sampletext):
    """
    Detect document language from filename using langid model
    Args:
        filename: The name of the file
    Returns:
        'de' for German, 'en' for English
    """
    try:
        name_without_ext = os.path.splitext(sampletext)[0]
        clean_name = name_without_ext.replace('_', ' ').replace('-', ' ')
        
        if len(clean_name.strip()) < 3:
            return "en"  
        
        detected_lang, confidence = langid.classify(clean_name)
        
        st.info(f"🔍 Detected language from filename '{sampletext}': {detected_lang.upper()} (confidence: {confidence:.2f})")
        
        if detected_lang == "de":
            return "de"
        else:
            return "en"  
            
    except Exception as e:
        st.warning(f"⚠️ Language detection failed for '{sampletext}': {str(e)}, defaulting to English")
        return "en"

def translate_text(text, target_lang):
    """Translate text to target language using deep-translator"""
    if not text or len(text.strip()) < 10:
        return text
    
    try:
        max_length = 4500
        if len(text) <= max_length:
            translator = GoogleTranslator(source='auto', target=target_lang)
            return translator.translate(text)
        else:
            paragraphs = text.split('\n\n')
            translated_paragraphs = []
            current_chunk = ""
            
            for para in paragraphs:
                if len(current_chunk) + len(para) < max_length:
                    current_chunk += para + "\n\n"
                else:
                    if current_chunk:
                        translator = GoogleTranslator(source='auto', target=target_lang)
                        translated_paragraphs.append(translator.translate(current_chunk))
                    current_chunk = para + "\n\n"
            
            if current_chunk:
                translator = GoogleTranslator(source='auto', target=target_lang)
                translated_paragraphs.append(translator.translate(current_chunk))
            
            return "\n\n".join(translated_paragraphs)
    except Exception as e:
        st.warning(f"Translation error: {str(e)}")
        return text

def get_file_hash(filepath):
    hash_md5 = hashlib.md5()
    with open(filepath, "rb") as f:
        for chunk in iter(lambda: f.read(4096), b""):
            hash_md5.update(chunk)
    return hash_md5.hexdigest()

def load_cache(cache_key):
    cache_file = os.path.join(CACHE_FOLDER, f"{cache_key}.pkl")
    if os.path.exists(cache_file):
        try:
            with open(cache_file, 'rb') as f:
                return pickle.load(f)
        except:
            return None
    return None

def save_cache(cache_key, data):
    cache_file = os.path.join(CACHE_FOLDER, f"{cache_key}.pkl")
    try:
        with open(cache_file, 'wb') as f:
            pickle.dump(data, f)
    except Exception as e:
        st.warning(f"⚠️ Cache save error: {str(e)}")

def clean_text(text):
    text = re.sub(r'\s+', ' ', text).strip()
    return text

def structure_text_into_paragraphs(text):
    if not text.strip():
        return ""
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    paragraphs = []
    current = []
    for line in lines:
        if re.match(r'^[\d]+[\.\)]\s|^[•\-\*]\s|^\n*📹', line):
            if current:
                paragraphs.append(' '.join(current))
                current = []
            paragraphs.append(line)
        else:
            current.append(line)
    if current:
        paragraphs.append(' '.join(current))
    return '\n\n'.join(paragraphs)

def create_smart_chunks(text, chunk_size=800, overlap=100, page_num=None, source_file=None, is_table=False, table_num=None, lang="en"):
    words = text.split()
    chunks = []
    metadata = {
        'page': str(page_num) if page_num is not None else "N/A",
        'source': source_file or "Unknown",
        'is_table': str(is_table),
        'table_number': str(table_num) if table_num else "N/A",
        'lang': lang
    }

    if len(words) <= chunk_size:
        if text.strip():
            chunks.append({'content': text.strip(), 'metadata': metadata})
        return chunks

    for i in range(0, len(words), chunk_size - overlap):
        chunk_words = words[i:i + chunk_size]
        chunk_text = " ".join(chunk_words)
        if len(chunk_words) >= 50:
            chunks.append({'content': chunk_text, 'metadata': metadata.copy()})
    return chunks

def format_table_as_structured_text(table, table_number=None):
    if not table or len(table) == 0:
        return ""
    headers = [str(cell).strip() or f"Col_{i+1}" for i, cell in enumerate(table[0])]
    text = f"\n📊 Table {table_number or ''}\n\n"
    text += "| " + " | ".join(headers) + " |\n"
    text += "| " + " --- |" * len(headers) + " |\n"
    for row in table[1:]:
        cells = [str(cell).strip() for cell in row]
        if any(cells):
            text += "| " + " | ".join(cells) + " |\n"
    return text

def extract_pdf_detailed(filepath):
    try:
        doc = fitz.open(filepath)
        if doc.is_encrypted and not doc.authenticate(PDF_PASSWORD):
            return None, "❌ Wrong PDF password"
    except Exception as e:
        return None, f"❌ PDF open error: {str(e)}"

    filename = os.path.basename(filepath)
    file_info = {'chunks': [], 'total_pages': len(doc), 'total_tables': 0}

    sample_text = ""
    for sample_page_num in range(min(3, len(doc))):  
        sample_page = doc[sample_page_num]
        sample_text += sample_page.get_text("text")
        if len(sample_text) > 1000:  
            break
    
    original_lang = detect_doc_language(sample_text)
    st.info(f"🔍 Detected language: {original_lang.upper()} for {filename}")

    for page_num in range(len(doc)):
        page = doc[page_num]

        text = page.get_text("text")
        if len(text.strip()) < 100:
            textpage = page.get_textpage_ocr(
                flags=fitz.TEXT_PRESERVE_LIGATURES | fitz.TEXT_PRESERVE_WHITESPACE,
                full=True,
                tessdata=TESSDATA_PATH
                
            )
            text = page.get_text("text", textpage=textpage)

        blocks = page.get_text("dict")["blocks"]
        page_text = f"# {filename} - Page {page_num + 1}\n\n"
        last_text_block = ""

        for block in blocks:
            if block.get("type") == 0:
                block_text = ""
                for line in block.get("lines", []):
                    for span in line.get("spans", []):
                        block_text += span.get("text", "")
                block_text = block_text.strip()

                if block_text:
                    structured = structure_text_into_paragraphs(block_text)
                    page_text += structured + "\n\n"
                    last_text_block = structured

        tables = page.find_tables()
        if tables:
            for t_num, table in enumerate(tables.tables, 1):
                file_info['total_tables'] += 1
                extracted = table.extract()

                if extracted:
                    table_text = format_table_as_structured_text(
                        extracted,
                        file_info['total_tables']
                    )

                    combined_text = ""
                    if last_text_block:
                        last_line = last_text_block.strip().split("\n")[-1].strip()
                        if (not last_line.endswith(".") or last_line.endswith(":") or len(last_line.split()) <= 12):
                            combined_text += last_text_block + "\n\n"
                    combined_text += table_text

                    table_chunks = create_smart_chunks(
                        combined_text,
                        chunk_size=1800,
                        overlap=0,
                        page_num=page_num + 1,
                        source_file=filename,
                        is_table=True,
                        table_num=file_info['total_tables'],
                        lang=original_lang
                    )
                    file_info['chunks'].extend(table_chunks)

                    target_lang = 'en' if original_lang == 'de' else 'de'
                    st.info(f"🔄 Translating table {file_info['total_tables']} to {target_lang.upper()}...")
                    translated_combined = translate_text(combined_text, target_lang)
                    
                    translated_table_chunks = create_smart_chunks(
                        translated_combined,
                        chunk_size=1800,
                        overlap=0,
                        page_num=page_num + 1,
                        source_file=f"{filename} ({target_lang.upper()})",
                        is_table=True,
                        table_num=file_info['total_tables'],
                        lang=target_lang
                    )
                    file_info['chunks'].extend(translated_table_chunks)

        page_chunks = create_smart_chunks(
            page_text,
            chunk_size=800,
            overlap=100,
            page_num=page_num + 1,
            source_file=filename,
            lang=original_lang
        )
        file_info['chunks'].extend(page_chunks)

        target_lang = 'en' if original_lang == 'de' else 'de'
        if page_num % 5 == 0:  
            st.info(f"🔄 Translating page {page_num + 1}/{len(doc)} to {target_lang.upper()}...")
        
        translated_page_text = translate_text(page_text, target_lang)
        translated_page_chunks = create_smart_chunks(
            translated_page_text,
            chunk_size=800,
            overlap=100,
            page_num=page_num + 1,
            source_file=f"{filename} ({target_lang.upper()})",
            lang=target_lang
        )
        file_info['chunks'].extend(translated_page_chunks)

    doc.close()
    return file_info, None

def extract_docx_detailed(filepath):
    doc = docx.Document(filepath)
    filename = os.path.basename(filepath)
    
    sample_text = ""
    for para in doc.paragraphs[:20]:  
        sample_text += para.text + " "
        if len(sample_text) > 1000:
            break
    
    original_lang = detect_doc_language(sample_text)
    st.info(f"🔍 Detected language: {original_lang.upper()} for {filename}")
    
    file_info = {
        'chunks': [],
        'total_pages': 1,
        'total_tables': 0,
        'pages_with_tables': [],
    }
   
    all_text = []
    table_counter = 0
   
    for element in doc.element.body:
        if element.tag.endswith('p'):
            for para in doc.paragraphs:
                if para._element == element:
                    text = clean_text(para.text)
                    if text:
                        structured = structure_text_into_paragraphs(text)
                        if structured:
                            all_text.append(structured)
                    break
       
        elif element.tag.endswith('tbl'):
            for table in doc.tables:
                if table._element == element:
                    file_info['total_tables'] += 1
                    table_counter += 1
                    table_text = format_table_as_structured_text(
                        [[cell.text for cell in row.cells] for row in table.rows],
                        table_counter
                    )
                    if table_text:
                        all_text.append(table_text)
                        
                        table_chunks = create_smart_chunks(
                            table_text,
                            chunk_size=2000,
                            overlap=0,
                            page_num=1,
                            source_file=filename,
                            is_table=True,
                            table_num=table_counter,
                            lang=original_lang
                        )
                        file_info['chunks'].extend(table_chunks)

                        target_lang = 'en' if original_lang == 'de' else 'de'
                        st.info(f"🔄 Translating table {table_counter} to {target_lang.upper()}...")
                        translated_table = translate_text(table_text, target_lang)
                        translated_table_chunks = create_smart_chunks(
                            translated_table,
                            chunk_size=2000,
                            overlap=0,
                            page_num=1,
                            source_file=f"{filename} ({target_lang.upper()})",
                            is_table=True,
                            table_num=table_counter,
                            lang=target_lang
                        )
                        file_info['chunks'].extend(translated_table_chunks)
                    break
   
    complete_text = "\n\n".join(all_text)
    
    text_chunks = create_smart_chunks(
        complete_text, 
        chunk_size=1500, 
        overlap=250,
        page_num=1,
        source_file=filename,
        lang=original_lang
    )
    file_info['chunks'].extend(text_chunks)

    target_lang = 'en' if original_lang == 'de' else 'de'
    st.info(f"🔄 Translating document to {target_lang.upper()}...")
    translated_complete = translate_text(complete_text, target_lang)
    translated_chunks = create_smart_chunks(
        translated_complete,
        chunk_size=1500,
        overlap=250,
        page_num=1,
        source_file=f"{filename} ({target_lang.upper()})",
        lang=target_lang
    )
    file_info['chunks'].extend(translated_chunks)
   
    if file_info['total_tables'] > 0:
        file_info['pages_with_tables'] = [1]
   
    return file_info, None

def extract_txt_detailed(filepath):
    filename = os.path.basename(filepath)
    
    with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
        text = f.read()
    
    original_lang = detect_doc_language(text)
    st.info(f"🔍 Detected language: {original_lang.upper()} for {filename}")
    
    structured_text = structure_text_into_paragraphs(text)
    
    chunks = create_smart_chunks(
        structured_text, 
        chunk_size=1500, 
        overlap=250,
        page_num=1,
        source_file=filename,
        lang=original_lang
    )
    
    file_info = {
        'chunks': chunks,
        'total_pages': 1,
        'total_tables': 0,
        'pages_with_tables': [],
    }

    target_lang = 'en' if original_lang == 'de' else 'de'
    st.info(f"🔄 Translating text file to {target_lang.upper()}...")
    translated_text = translate_text(structured_text, target_lang)
    translated_chunks = create_smart_chunks(
        translated_text,
        chunk_size=1500,
        overlap=250,
        page_num=1,
        source_file=f"{filename} ({target_lang.upper()})",
        lang=target_lang
    )
    file_info['chunks'].extend(translated_chunks)
    
    return file_info, None
    
    
def get_files_from_folder():
    return glob.glob(os.path.join(DOCS_FOLDER, "*.[pP][dD][fF]")) + \
           glob.glob(os.path.join(DOCS_FOLDER, "*.[dD][oO][cC][xX]")) + \
           glob.glob(os.path.join(DOCS_FOLDER, "*.txt"))



